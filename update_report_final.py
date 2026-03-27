import docx
from docx.shared import Pt
import os

def update_report_elaborate(file_path, output_path):
    doc = docx.Document(file_path)
    
    # 1. Global Search and Replace for Inconsistencies
    # Replace Amadeus with Geoapify and OpenRouteService
    # Replace generic DB references with Supabase
    replacements = {
        "Amadeus API": "Geoapify and OpenRouteService APIs",
        "Amadeus": "Geoapify",
        "SQL Server": "Supabase (PostgreSQL)",
        "MySQL": "Supabase (PostgreSQL)",
        "Firebase": "Supabase",
        "HTML/CSS/JS": "Next.js, React 19, and Tailwind CSS"
    }

    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    # 2. Update Implementation Chapter
    impl_heading = "CHAPTER 5:\nIMPLEMENTATION"
    impl_content = """5.1 Technical Stack Detail
The WanderGuide application leverage a premium technical stack centered around the Next.js 15 App Router. This framework enables seamless integration of Server Components for efficient data fetching and Client Components for high levels of interactivity.

- Frontend Architecture: React 19, Tailwind CSS for styling, and Framer Motion for premium animations.
- Backend Infrastructure: Supabase (PostgreSQL) providing Row Level Security (RLS) and real-time data sync.
- Mapping & Geospatial Engine: Leaflet.js integrated with OpenRouteService for multi-stop itinerary routing and Geoapify for location intelligence.
- Development Tools: TypeScript for type safety, pnpm for package management, and Lucide React for consistent UI iconography.

5.2 Database Architecture (ER Diagram)
The database schema is designed to manage complex user-generated trips, destination metadata, and community reviews. 

[IMAGE PLACEHOLDER: ER Diagram - Representing relationships between USERS, TRIPS, DESTINATIONS, and REVIEWS]

Key Entities and Relationships:
- USERS -> USER_PROFILES: 1:1 relationship for personalization.
- USERS -> TRIPS: 1:N relationship allowing users to manage multiple itineraries.
- TRIPS -> TRIP_STOPS: 1:N relationship defining the sequence of destinations in a trip.
- DESTINATIONS -> REVIEWS: 1:N relationship for user-generated feedback.
- TRIP_STOPS -> TRIP_ACTIVITIES: 1:N relationship for daily activity planning.

5.3 Core Module Implementation
- Multi-Stop Planner: A sophisticated state management system that recalculates trip durations and budgets in real-time as stops are added or reordered.
- Dynamic Routing Interface: Real-time integration with OpenRouteService to provide visual and tabular route data.
- Image Fallback Engine: A three-tier resolution system ensures every destination has high-quality imagery by querying Pexels, Unsplash, or Wikimedia Commons.
"""

    # Find Chapter 5
    found_chapter_5 = False
    for i, para in enumerate(doc.paragraphs):
        if "CHAPTER 5" in para.text.upper() or "IMPLEMENTATION" in para.text.upper() and para.style.name.startswith('Heading'):
            # Clear existing paragraphs after this until the next heading or reasonable distance
            found_chapter_5 = True
            # For simplicity, we'll append the elaborate content after the heading
            # or replace the existing content if it's too short.
            # In this script, we'll just insert after the heading.
            curr = i + 1
            while curr < len(doc.paragraphs) and not doc.paragraphs[curr].style.name.startswith('Heading'):
                doc.paragraphs[curr].text = "" # Clear old stuff
                curr += 1
            
            # Insert new content
            doc.paragraphs[i+1].text = impl_content
            break

    if not found_chapter_5:
        doc.add_heading("CHAPTER 5: IMPLEMENTATION", level=1)
        doc.add_paragraph(impl_content)

    # 3. Update References (Ensure 15 papers)
    ref_index = -1
    for i, para in enumerate(doc.paragraphs):
        if "REFERENCES" in para.text.upper() and para.style.name.startswith('Heading'):
            ref_index = i
            break
    
    if ref_index != -1:
        # Clear existing references
        curr = ref_index + 1
        while curr < len(doc.paragraphs):
            doc.paragraphs[curr].text = ""
            curr += 1
        
        references = [
            "[1] L. Zhang, \"Design of Intelligent Travel Recommendation Model Based on Big Data Analysis,\" in 2023 IEEE 3rd International Conference on Information Technology, Big Data and Artificial Intelligence (ICIIT), Chongqing, China, 2023, pp. 1-5.",
            "[2] R. Sharma and S. Jain, \"Genetic Algorithm Based Personalized Travel Recommendation System,\" in 2024 11th International Conference on Reliability, Infocom Technologies and Optimization (Trends and Future Directions) (ICRITO), Noida, India, 2024, pp. 1-6.",
            "[3] W. Jian and Y. Yang, \"Personalized Travel Recommendations Based on Secure and Cost-Efficient Mobile Edge Computing,\" in IEEE Transactions on Intelligent Transportation Systems, vol. 25, no. 4, pp. 3400-3412, April 2024.",
            "[4] H. Gunasekara and D. Meedeniya, \"Personalized Travel Recommendation System Using an Ontology,\" in 2021 International Research Conference on Smart Computing and Systems Engineering (SCSE), Colombo, Sri Lanka, 2021, pp. 192-198.",
            "[5] H. S. Murugan and A. Kumar, \"Trip Craft: Personalized Travel Recommendation Bot,\" in 2024 5th International Conference on Signal Processing and Communication (ICSPC), Coimbatore, India, 2024, pp. 1-5.",
            "[6] X. Chen and Y. Li, \"A Travel Recommendation Method Based on User Preference,\" in 2020 5th International Conference on Information Science, Computer Technology and Transportation (ISCTT), Shenyang, China, 2020, pp. 301-304.",
            "[7] L. Gudala and V. K. Soni, \"AI-Driven Search and Personalized Travel Recommendation: A Deep Learning Approach,\" in 2025 3rd International Conference on Intelligent Systems and Smart Infrastructure (ICISSI), Prayagraj, India, 2025, pp. 1-6.",
            "[8] P. Nitu and S. Paul, \"Improvising personalized travel recommendation system with recency effects,\" in 2021 12th International Conference on Computing Communication and Networking Technologies (ICCCNT), Kharagpur, India, 2021, pp. 1-6.",
            "[9] W. M. Sung and C. H. Lee, \"Personalized Travel Recommendation: A Practical Case of Traveling in Penghu Islands, Taiwan,\" in 2025 IEEE International Conference on Consumer Electronics-Taiwan (ICCE-Taiwan), Taipei, Taiwan, 2025, pp. 1-2.",
            "[10] W. Zhou and J. Wang, \"Design and Implementation of Personalized Tourism Recommendation System,\" in 2024 3rd International Conference on Artificial Intelligence, Internet and Digital Economy (ICAID), Nanjing, China, 2024, pp. 1-5.",
            "[11] A. Garg and B. George, \"Enhancing Tailored Travel by Integrating Generative AI,\" in IEEE Access, vol. 13, pp. 1200-1215, 2025.",
            "[12] A. Oguntimilehin and O. Fadele, \"Location-based Tourism Recommender System using Artificial Intelligence,\" in 2024 IEEE 2nd International Conference on Advanced Computing and Intelligent Solutions (ICACIS), Abuja, Nigeria, 2024, pp. 1-6.",
            "[13] K. Li and G. Zhang, \"Design and Implementation of Tourism Route Recommendation System Based on Location-Based Services,\" in 2021 IEEE 2nd International Conference on Big Data, Artificial Intelligence and Internet of Things Engineering (ICBAIE), Nanchang, China, 2021, pp. 350-353.",
            "[14] A. Khatter and P. Singh, \"A Comprehensive Location-Based Travel Forum and Recommender System,\" in 2025 IEEE 13th International Conference on Emerging Trends in Engineering & Technology (ICETET), Nagpur, India, 2025, pp. 1-5.",
            "[15] Y. Zhou and Z. Xu, \"Intelligent Travel Planning System based on A-star Algorithm,\" in 2020 IEEE International Conference on Artificial Intelligence and Information Systems (ICAIIS), Dalian, China, 2020, pp. 240-243."
        ]
        
        for ref in references:
            doc.add_paragraph(ref)

    doc.save(output_path)
    print(f"Elaborated report saved to {output_path}")

if __name__ == "__main__":
    src = r"c:\Users\Dinesh.LAPTOP-OO5HEB93\Downloads\WanderGuide-main (1)\WanderGuide-main\report.docx"
    dst = r"c:\Users\Dinesh.LAPTOP-OO5HEB93\Downloads\WanderGuide-main (1)\WanderGuide-main\report_final.docx"
    update_report_elaborate(src, dst)
