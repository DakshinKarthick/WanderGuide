import docx
from docx.shared import Pt
import os

def update_report(file_path, output_path):
    doc = docx.Document(file_path)
    
    # Content to insert
    impl_content = """### 1. System Architecture
The WanderGuide application is developed using the Next.js App Router (version 15+), which provides a robust framework for building fast, SEO-friendly, and scalable web applications. The architecture follows a full-stack approach where the frontend and backend are seamlessly integrated through Next.js API routes and Server Components.

### 2. Technical Stack
- Frontend: Built with React 19 and Tailwind CSS. Framer Motion is used for smooth page transitions and micro-interactions, while Lucide React provides a consistent iconography set.
- Backend & Database: Supabase (PostgreSQL) is employed for real-time data storage, user authentication, and managing relational data such as destinations, reviews, and notifications.
- Mapping Services: The project integrates Leaflet.js for interactive map rendering and OpenRouteService for multi-stop routing and distance estimations.
- Third-Party APIs: 
    - Geoapify: Used for location search and geocoding.
    - Pexels / Unsplash: Dynamic image resolution for destinations.

### 3. Key Implementation Details
- Multi-Stop Planner: A specialized module that manages the state of a user's trip, allowing for dynamic addition, removal, and reordering of destinations with real-time route updates.
- Pricing Engine: Integrates distance data with predefined pricing models to provide cost estimates for transport, accommodation, and daily activities.
- Image Fallback System: To ensure no destination is left without a visual, the system implements a tiered fallback: Pexels -> Unsplash -> Wikimedia Commons.

[IMAGE PLACEHOLDER: WanderGuide Landing Page - Showing Hero section and Search Bar]

[IMAGE PLACEHOLDER: Multi-Stop Planner Interface - Showing destination list and dynamic map]

[IMAGE PLACEHOLDER: Destination Details Modal - Showing activities, pricing, and reviews]

[IMAGE PLACEHOLDER: Pricing Summary - Showing cost breakdown for various trip categories]
"""

    references = [
        "[1] L. Zhang, \"Design of Intelligent Travel Recommendation Model Based on Big Data Analysis,\" in 2023 IEEE 3rd International Conference on Information Technology, Big Data and Artificial Intelligence (ICIIT), Chongqing, China, 2023, pp. 1-5.",
        "[2] R. Sharma and S. Jain, \"Genetic Algorithm Based Personalized Travel Recommendation System,\" in 2024 11th International Conference on Reliability, Infocom Technologies and Optimization (Trends and Future Directions) (ICRITO), Noida, India, 2024, pp. 1-6.",
        "[3] W. Jian and Y. Yang, \"Personalized Travel Recommendations Based on Secure and Cost-Efficient Mobile Edge Computing,\" in IEEE Transactions on Intelligent Transportation Systems, vol. 25, no. 4, pp. 3400-3412, April 2024.",
        "[4] H. Gunasekara and D. Meedeniya, \"Personalized Travel Recommendation System Using an Ontology,\" in 2021 International Research Conference on Smart Computing and Systems Engineering (SCSE), Colombo, Sri Lanka, 2021, pp. 192-198.",
        "[5] H. S. Murugan and A. Kumar, \"Trip Craft: Personalized Travel Recommendation Bot,\" in 2024 5th International Conference on Signal Processing and Communication (ICSPC), Coimbatore, India, 2024, pp. 1-5.",
        "[6] X. Chen and Y. Li, \"A Travel Recommendation Method Based on User Preference,\" in 2020 5th International Conference on Information Science, Computer Technology and Transportation (ISCTT), Shenyang, China, 2020, pp. 301-304.",
        "[7] L. Gudala and V. K. Soni, \"AI-Driven Search and Personalized Travel Recommendation: A Deep Learning Approach,\" in 2025 3rd International Conference on Intelligent Systems and Smart Infrastructure (ICISSI), Prayagraj, India, 2025, pp. 1-6.",
        "[8] P. Nitu and S. Paul, \"Improvising personalized travel recommendation system with recency effects,\" in 2021 12th International Conference on Computing Communication and Networking Technologies (ICCCNT), Kharagpur, India, 2021, pp. 1-6.",
        "[9] W. M. Sung and C. H. Lee, \"Personalized Travel Recommendation: A Practical Case of Traveling in Penghu Islands, Taiwan,\" in 2025 IEEE International Conference on Consumer Electronics-Taiwan (ICCE-Taiwan), Taipei, Taiwan, 2025, pp. 1-2.",
        "[10] W. Zhou and J. Wang, \"Design and Implementation of Personalized Tourism Recommendation System,\" in 2024 3rd International Conference on Artificial Intelligence, Internet and Digital Economy (ICAID), Nanjing, China, 2024, pp. 1-5.",
        "[11] A. Garg and B. George, \"Enhancing Tailored Travel by Integrating Generative AI,\" in IEEE Access, vol. 13, pp. 1200-1215, 2025.",
        "[12] A. Oguntimilehin and O. Fadele, \"Location-based Tourism Recommender System using Artificial Intelligence,\" in 2024 IEEE 2nd International Conference on Advanced Computing and Intelligent Solutions (ICACIS), Abuja, Nigeria, 2024, pp. 1-6.",
        "[13] K. Li and G. Zhang, \"Design and Implementation of Tourism Route Recommendation System Based on Location-Based Services,\" in 2021 IEEE 2nd International Conference on Big Data, Artificial Intelligence and Internet of Things Engineering (ICBAIE), Nanchang, China, 2021, pp. 350-353.",
        "[14] A. Khatter and P. Singh, \"A Comprehensive Location-Based Travel Forum and Recommender System,\" in 2025 IEEE 13th International Conference on Emerging Trends in Engineering & Technology (ICETET), Nagpur, India, 2025, pp. 1-5.",
        "[15] Y. Zhou and Z. Xu, \"Intelligent Travel Planning System based on A-star Algorithm,\" in 2020 IEEE International Conference on Artificial Intelligence and Information Systems (ICAIIS), Dalian, China, 2020, pp. 240-243."
    ]

    # Find where to insert Implementation
    impl_index = -1
    ref_index = -1
    
    for i, para in enumerate(doc.paragraphs):
        text_upper = para.text.upper()
        if "IMPLEMENTATION" in text_upper and para.style.name.startswith('Heading'):
            impl_index = i
        if "REFERENCES" in text_upper and para.style.name.startswith('Heading'):
            ref_index = i

    # Insert Implementation details
    if impl_index != -1:
        # Insert after the Implementation heading
        # Note: python-docx doesn't easily insert after an index without re-copying normally,
        # but we can add paragraphs and then reorder or just append if it's the right place.
        # Here we'll just append to the end and note it, or try to insert if possible.
        pass
    
    # Simple approach: Create a new doc with these sections appended or inserted
    # Since we can't easily "insert" in the middle of list of paragraphs in python-docx easily,
    # we will append a NEW section at the end if we can't find a clear spot, 
    # but let's try to find the "REFERENCES" and insert BEFORE it.
    
    if ref_index != -1:
        # Insert Implementation BEFORE References
        doc.paragraphs[ref_index].insert_paragraph_before("IMPLEMENTATION", style='Heading 1')
        new_para = doc.paragraphs[ref_index] # This is the newly inserted "IMPLEMENTATION"
        for line in impl_content.split('\n'):
            doc.paragraphs[ref_index+1].insert_paragraph_before(line)
        
        # Add References at the end of the document or after the References heading
        for ref in references:
            doc.add_paragraph(ref)
    else:
        # Just append everything
        doc.add_heading("IMPLEMENTATION", level=1)
        doc.add_paragraph(impl_content)
        doc.add_heading("REFERENCES", level=1)
        for ref in references:
            doc.add_paragraph(ref)

    doc.save(output_path)
    print(f"Updated report saved to {output_path}")

if __name__ == "__main__":
    src = r"c:\Users\Dinesh.LAPTOP-OO5HEB93\Downloads\WanderGuide-main (1)\WanderGuide-main\report.docx"
    dst = r"c:\Users\Dinesh.LAPTOP-OO5HEB93\Downloads\WanderGuide-main (1)\WanderGuide-main\report_updated.docx"
    update_report(src, dst)
